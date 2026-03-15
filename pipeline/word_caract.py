"""
word_caract.py
Genera informe Word de caracterización al ingreso (TOP1).
Reemplaza pdf_caract.py — misma estructura, mismo contenido, formato Word.
"""
import os, io, unicodedata
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Rutas (inyectadas por runner) ─────────────────────────────────────────────
INPUT_FILE  = None
OUTPUT_FILE = None
SHEET_NAME  = 'Base Wide'
FILTRO_CENTRO = None
NOMBRE_SERVICIO = 'Sistema de Monitoreo'
PERIODO = ''

# ── Colores ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x38, 0x64)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
ACCENT = RGBColor(0x00, 0xB0, 0xF0)
ORANGE = RGBColor(0xC8, 0x59, 0x0A)
LIGHT  = RGBColor(0xEE, 0xF4, 0xFB)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x55, 0x55, 0x55)
GREEN  = RGBColor(0x53, 0x81, 0x35)

C_NAVY   = '1F3864'; C_BLUE = '2E75B6'; C_ACCENT = '00B0F0'
C_LIGHT  = 'EEF4FB'; C_WHITE = 'FFFFFF'; C_ORANGE = 'C8590A'

MPL_NAVY   = '#1F3864'; MPL_BLUE = '#2E75B6'; MPL_ACCENT = '#00B0F0'
MPL_ORANGE = '#C8590A'; MPL_GREEN = '#538135'; MPL_GRAY = '#888888'

def _norm(s):
    return unicodedata.normalize('NFD', str(s).lower()).encode('ascii','ignore').decode()

# ── Helpers Word ──────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, hex_color='DDDDDD'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), hex_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(13)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = BLUE
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
    return p

def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'Nota: {text}')
    run.font.name = 'Calibri'
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = GRAY
    return p

def fig_to_docx_image(fig, width_cm=14):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf, Cm(width_cm)

def add_section_header(doc, num, title):
    """Agrega un encabezado de sección con fondo azul navy."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, C_NAVY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f'  {num}. {title.upper()}')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = WHITE
    doc.add_paragraph()

def add_kpi_row(doc, kpis):
    """Fila de KPIs: lista de (valor, label)."""
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (val, lbl) in enumerate(kpis):
        # Valor
        c_val = table.rows[0].cells[j]
        set_cell_bg(c_val, C_LIGHT)
        p = c_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.name = 'Calibri'; run.font.size = Pt(24)
        run.font.bold = True; run.font.color.rgb = BLUE
        # Label
        c_lbl = table.rows[1].cells[j]
        set_cell_bg(c_lbl, C_LIGHT)
        p2 = c_lbl.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(lbl)
        run2.font.name = 'Calibri'; run2.font.size = Pt(8.5)
        run2.font.color.rgb = GRAY
    doc.add_paragraph()

# ── Gráficos ──────────────────────────────────────────────────────────────────
def _ax_style(ax):
    ax.set_facecolor('white')
    ax.spines[['top','right']].set_visible(False)
    ax.tick_params(labelsize=8)

def g_sexo(R):
    fig, ax = plt.subplots(figsize=(5, 3))
    labels = ['Hombres', 'Mujeres']
    vals   = [R.get('n_hombre',0), R.get('n_mujer',0)]
    colors = [MPL_BLUE, MPL_ACCENT]
    bars = ax.bar(labels, vals, color=colors, width=0.5)
    for b, v in zip(bars, vals):
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.3,
                str(v), ha='center', va='bottom', fontsize=10, fontweight='bold')
    _ax_style(ax); ax.yaxis.set_visible(False)
    ax.set_title('Distribución por Sexo', fontsize=10, fontweight='bold', color=MPL_NAVY)
    plt.tight_layout()
    return fig

def g_edad(R):
    fig, ax = plt.subplots(figsize=(6, 3))
    grupos  = R.get('edad_grupos', {})
    if grupos:
        ax.bar(list(grupos.keys()), list(grupos.values()),
               color=MPL_BLUE, width=0.6)
    _ax_style(ax)
    ax.set_title('Distribución por Rango de Edad', fontsize=10, fontweight='bold', color=MPL_NAVY)
    plt.tight_layout()
    return fig

def g_sust(R):
    fig, ax = plt.subplots(figsize=(5, 4))
    sust = R.get('sust_dist', {})
    if sust:
        sd = pd.DataFrame(list(sust.items()), columns=['S','n']).sort_values('n')
        tot = sd['n'].sum()
        colors = [MPL_BLUE if i%2==0 else MPL_ACCENT for i in range(len(sd))]
        ax.barh(sd['S'], sd['n'], color=colors, height=0.6)
        for b, v in zip(ax.patches, sd['n']):
            pct = round(v/tot*100,1) if tot else 0
            ax.text(b.get_width()+0.2, b.get_y()+b.get_height()/2,
                    f'{v} ({pct}%)', va='center', fontsize=8)
        ax.spines[['top','right','bottom']].set_visible(False)
        ax.xaxis.set_visible(False)
    ax.set_title('Sustancia Principal', fontsize=10, fontweight='bold', color=MPL_NAVY)
    plt.tight_layout()
    return fig

def g_transgresion(R):
    fig, ax = plt.subplots(figsize=(5, 3))
    n_trans = R.get('n_transgresores', 0)
    n_no    = R.get('N', 0) - n_trans
    if n_trans + n_no > 0:
        wedges, _, autotexts = ax.pie(
            [n_trans, n_no],
            labels=[f'Con transgresión\n({n_trans})', f'Sin transgresión\n({n_no})'],
            colors=[MPL_ORANGE, MPL_BLUE],
            autopct='%1.0f%%', startangle=90,
            wedgeprops={'edgecolor':'white','linewidth':1.5},
            textprops={'fontsize':8}
        )
        for at in autotexts: at.set_color('white'); at.set_fontweight('bold')
    ax.set_title('Transgresión a la Norma Social', fontsize=10, fontweight='bold', color=MPL_NAVY)
    plt.tight_layout()
    return fig

def g_tipos_transgresion(R):
    fig, ax = plt.subplots(figsize=(6, 3.5))
    tipos = R.get('transgresion_tipos', {})
    if tipos:
        td = pd.DataFrame(list(tipos.items()), columns=['T','n']).sort_values('n')
        ax.barh(td['T'], td['n'], color=MPL_ORANGE, height=0.6)
        ax.spines[['top','right','bottom']].set_visible(False)
        ax.xaxis.set_visible(False)
        for b, v in zip(ax.patches, td['n']):
            ax.text(b.get_width()+0.1, b.get_y()+b.get_height()/2,
                    str(v), va='center', fontsize=8)
    ax.set_title('Tipos de Transgresión', fontsize=10, fontweight='bold', color=MPL_NAVY)
    plt.tight_layout()
    return fig

def g_salud(R):
    fig, ax = plt.subplots(figsize=(5, 3))
    sal = R.get('salud_dist', {})
    if sal:
        labels = list(sal.keys())
        vals   = list(sal.values())
        colors = [MPL_GREEN, MPL_BLUE, MPL_ACCENT, MPL_ORANGE, '#C00000'][:len(labels)]
        bars = ax.bar(labels, vals, color=colors[:len(labels)], width=0.6)
        for b, v in zip(bars, vals):
            ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.2,
                    str(v), ha='center', va='bottom', fontsize=8)
        ax.yaxis.set_visible(False)
    _ax_style(ax)
    ax.set_title('Autopercepción del Estado de Salud', fontsize=10, fontweight='bold', color=MPL_NAVY)
    plt.tight_layout()
    return fig

# ── Carga de datos ────────────────────────────────────────────────────────────
def cargar_datos():
    df = pd.read_excel(INPUT_FILE, sheet_name=SHEET_NAME, header=1)
    df.columns = [str(c) for c in df.columns]

    N  = len(df)
    R  = {'N': N, 'd': df}

    # Detectar columnas
    def _col(*keys):
        for c in df.columns:
            nc = _norm(c)
            if any(k in nc for k in keys): return c
        return None

    col_sexo  = _col('sexo','genero','género')
    col_fn    = _col('fecha de nacimiento','fecha_nac','nacimiento')
    col_fecha = _col('fecha entrevista_top1','fecha entrevista')
    col_sust  = next((c for c in df.columns
                      if any(k in _norm(c) for k in
                             ['sustancia principal','cual considera','genera mas problemas'])
                      and '_TOP1' in c and 'RAW' not in c), None)
    col_trans = _col('transgresion','transgresión')
    col_salud = _col('estado de salud','autopercepcion','autopercepción')
    col_edad  = _col('rango de edad','rango_edad','grupo edad')
    col_viv   = _col('vivienda','condicion de vivienda')

    # Sexo
    if col_sexo:
        sc = df[col_sexo].astype(str).str.strip().str.upper()
        nv = int(sc.isin(['H','M']).sum())
        R['n_hombre']   = int((sc=='H').sum())
        R['n_mujer']    = int((sc=='M').sum())
        R['nv_sex']     = nv
        R['pct_hombre'] = round(R['n_hombre']/nv*100,1) if nv>0 else 0
        R['pct_mujer']  = round(R['n_mujer'] /nv*100,1) if nv>0 else 0
    else:
        R.update({'n_hombre':0,'n_mujer':0,'nv_sex':0,'pct_hombre':0,'pct_mujer':0})

    # Edad
    R['edad_media'] = 0; R['edad_grupos'] = {}
    if col_fn and col_fecha:
        fn  = pd.to_datetime(df[col_fn], errors='coerce')
        ref = pd.to_datetime(df[col_fecha], errors='coerce').fillna(pd.Timestamp.now())
        edad = ((ref-fn).dt.days/365.25).round(1)
        edad = edad[(edad>=10)&(edad<=100)]
        R['edad_media'] = round(float(edad.mean()),1) if len(edad) else 0
        bins = [0,17,24,34,44,54,64,200]
        labs = ['<18','18-24','25-34','35-44','45-54','55-64','65+']
        cut  = pd.cut(edad, bins=bins, labels=labs)
        R['edad_grupos'] = cut.value_counts().sort_index().to_dict()
    elif col_edad:
        R['edad_grupos'] = df[col_edad].value_counts().head(8).to_dict()

    # Sustancia principal
    R['sust_dist'] = {}; R['sust_top1'] = ''; R['sust_top1_pct'] = 0
    if col_sust:
        vc = df[col_sust].dropna().value_counts()
        R['sust_dist']    = vc.head(8).to_dict()
        if len(vc):
            R['sust_top1']     = vc.index[0]
            R['sust_top1_pct'] = round(vc.iloc[0]/N*100,1)

    # Transgresión
    R['n_transgresores'] = 0; R['pct_transgresores'] = 0; R['transgresion_tipos'] = {}
    if col_trans:
        trans = df[col_trans].astype(str).str.strip()
        R['n_transgresores']    = int((trans=='Sí').sum())
        R['pct_transgresores']  = round(R['n_transgresores']/N*100,1) if N>0 else 0
        # Tipos de transgresión
        tipo_cols = [c for c in df.columns if any(k in _norm(c)
                     for k in ['tipo de transgresion','tipo_transgresion'])]
        if tipo_cols:
            tipos_serie = df[tipo_cols[0]].dropna()
            R['transgresion_tipos'] = tipos_serie.value_counts().head(6).to_dict()

    # Salud
    R['salud_media'] = 0; R['salud_dist'] = {}
    if col_salud:
        sal = pd.to_numeric(df[col_salud], errors='coerce').dropna()
        R['salud_media'] = round(float(sal.mean()),1) if len(sal) else 0
        bins = [0,4,8,12,16,20]
        labs = ['0-4','5-8','9-12','13-16','17-20']
        cut  = pd.cut(sal, bins=bins, labels=labs, include_lowest=True)
        R['salud_dist'] = cut.value_counts().sort_index().to_dict()

    # Vivienda
    R['vivienda_dist'] = {}
    if col_viv:
        R['vivienda_dist'] = df[col_viv].value_counts().head(6).to_dict()

    # Período
    global PERIODO
    if col_fecha:
        fechas = pd.to_datetime(df[col_fecha], errors='coerce').dropna()
        if len(fechas):
            MESES = {1:'Enero',2:'Febrero',3:'Marzo',4:'Abril',5:'Mayo',6:'Junio',
                     7:'Julio',8:'Agosto',9:'Septiembre',10:'Octubre',11:'Noviembre',12:'Diciembre'}
            f1, f2 = fechas.min(), fechas.max()
            if f1.year==f2.year and f1.month==f2.month:
                PERIODO = f'{MESES[f1.month]} {f1.year}'
            elif f1.year==f2.year:
                PERIODO = f'{MESES[f1.month]}–{MESES[f2.month]} {f1.year}'
            else:
                PERIODO = f'{MESES[f1.month]} {f1.year} – {MESES[f2.month]} {f2.year}'

    return R

# ── Construir documento Word ──────────────────────────────────────────────────
def build_word(R):
    doc = Document()

    # Márgenes
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── Portada ────────────────────────────────────────────────────────────────
    t_port = doc.add_table(rows=1, cols=1)
    t_port.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t_port.rows[0].cells[0]
    set_cell_bg(c, C_NAVY)
    for txt, sz, bold in [
        ('INFORME DE CARACTERIZACIÓN', 18, True),
        ('Monitoreo de Resultados de Tratamiento — Instrumento TOP', 11, False),
        (NOMBRE_SERVICIO.upper(), 14, True),
        (PERIODO, 10, False),
    ]:
        p = c.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.font.name = 'Calibri'; run.font.size = Pt(sz)
        run.font.bold = bold; run.font.color.rgb = WHITE
    doc.add_paragraph()

    # ── Presentación ──────────────────────────────────────────────────────────
    add_section_header(doc, '', 'Presentación')
    add_body(doc,
        f'Este informe presenta los resultados de caracterización al ingreso de {R["N"]} '
        f'personas que iniciaron tratamiento por consumo de sustancias psicoactivas. '
        f'Los datos fueron recopilados mediante el instrumento TOP (Treatment Outcomes Profile) '
        f'durante el período {PERIODO}.')
    doc.add_paragraph()

    # KPIs
    kpis = [(R['N'], 'Personas ingresaron')]
    if R.get('pct_hombre'): kpis.append((f'{R["pct_hombre"]}%', 'Son hombres'))
    if R.get('edad_media'):  kpis.append((R['edad_media'], 'Edad promedio'))
    if R.get('sust_top1'):   kpis.append((R['sust_top1'], 'Sustancia principal'))
    add_kpi_row(doc, kpis)

    # ── 1. Antecedentes Generales ──────────────────────────────────────────────
    add_section_header(doc, '1', 'Antecedentes Generales')

    add_heading(doc, '1.1. Distribución de Personas según Sexo', level=2)
    if R.get('n_hombre',0) + R.get('n_mujer',0) > 0:
        fig = g_sexo(R)
        buf, w = fig_to_docx_image(fig, 12)
        doc.add_picture(buf, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_body(doc,
            f'Del total de {R["N"]} personas que ingresaron a tratamiento, '
            f'{R["n_hombre"]} ({R["pct_hombre"]}%) son hombres y '
            f'{R["n_mujer"]} ({R["pct_mujer"]}%) son mujeres.')

    add_heading(doc, '1.2. Distribución de Personas según Edad', level=2)
    if R.get('edad_grupos'):
        fig = g_edad(R)
        buf, w = fig_to_docx_image(fig, 13)
        doc.add_picture(buf, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if R.get('edad_media'):
            add_body(doc, f'La edad promedio de las personas al ingreso es de {R["edad_media"]} años.')

    # ── 2. Consumo de Sustancias ───────────────────────────────────────────────
    add_section_header(doc, '2', 'Consumo de Sustancias')

    add_heading(doc, '2.1. Sustancia Principal al Ingreso', level=2)
    if R.get('sust_dist'):
        fig = g_sust(R)
        buf, w = fig_to_docx_image(fig, 12)
        doc.add_picture(buf, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_body(doc,
            f'La sustancia principal más frecuente al ingreso es {R["sust_top1"]} '
            f'({R["sust_top1_pct"]}% de las personas).')

    # ── 3. Transgresión a la Norma Social ─────────────────────────────────────
    add_section_header(doc, '3', 'Transgresión a la Norma Social')

    add_heading(doc, '3.1. Transgresión a la Norma Social', level=2)
    if R['N'] > 0:
        fig = g_transgresion(R)
        buf, w = fig_to_docx_image(fig, 11)
        doc.add_picture(buf, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_body(doc,
            f'Un total de {R["n_transgresores"]} personas ({R["pct_transgresores"]}%) '
            f'reportaron haber cometido algún tipo de transgresión a la norma social '
            f'en las cuatro semanas previas al ingreso a tratamiento.')

    if R.get('transgresion_tipos'):
        add_heading(doc, '3.2. Distribución por Tipo de Transgresión', level=2)
        fig = g_tipos_transgresion(R)
        buf, w = fig_to_docx_image(fig, 13)
        doc.add_picture(buf, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 4. Salud, Calidad de Vida y Vivienda ──────────────────────────────────
    add_section_header(doc, '4', 'Salud, Calidad de Vida y Vivienda')

    add_heading(doc, '4.1. Autopercepción del Estado de Salud', level=2)
    if R.get('salud_dist'):
        fig = g_salud(R)
        buf, w = fig_to_docx_image(fig, 12)
        doc.add_picture(buf, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_body(doc,
            f'La puntuación promedio de salud autopercibida al ingreso es de '
            f'{R["salud_media"]} puntos (escala 0–20).')

    if R.get('vivienda_dist'):
        add_heading(doc, '4.2. Condiciones de Vivienda al Ingreso', level=2)
        # Tabla de vivienda
        items = list(R['vivienda_dist'].items())
        tbl = doc.add_table(rows=len(items)+1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for j, hdr in enumerate(['Condición de vivienda','N']):
            c = tbl.rows[0].cells[j]
            set_cell_bg(c, C_NAVY)
            p = c.paragraphs[0]
            run = p.add_run(hdr)
            run.font.bold = True; run.font.size = Pt(9.5)
            run.font.color.rgb = WHITE; run.font.name = 'Calibri'
        for i, (lbl, val) in enumerate(items):
            bg = C_LIGHT if i%2==0 else C_WHITE
            for j, txt in enumerate([str(lbl), str(val)]):
                c = tbl.rows[i+1].cells[j]
                set_cell_bg(c, bg)
                p = c.paragraphs[0]
                run = p.add_run(txt)
                run.font.size = Pt(9.5); run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0x33,0x33,0x33)

    # ── Pie de página ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'Informe generado automáticamente por QALAT · Sistema de Monitoreo TOP · {PERIODO}')
    run.font.size = Pt(8); run.font.italic = True
    run.font.color.rgb = GRAY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_FILE)
    print(f'  ✓ Word generado: {OUTPUT_FILE}')


if __name__ == '__main__':
    R = cargar_datos()
    build_word(R)
