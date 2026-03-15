"""
word_seg.py
Genera informe Word de seguimiento TOP1 vs TOP2.
Reemplaza pdf_seg.py — mismo contenido, formato Word.
"""
import os, unicodedata
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_FILE  = None
OUTPUT_FILE = None
SHEET_NAME  = 'Base Wide'
FILTRO_CENTRO = None
NOMBRE_SERVICIO = 'Sistema de Monitoreo'
PERIODO = ''

NAVY  = RGBColor(0x1F,0x38,0x64); BLUE  = RGBColor(0x2E,0x75,0xB6)
ACCENT= RGBColor(0x00,0xB0,0xF0); ORANGE= RGBColor(0xC8,0x59,0x0A)
WHITE = RGBColor(0xFF,0xFF,0xFF); GRAY  = RGBColor(0x55,0x55,0x55)
GREEN = RGBColor(0x53,0x81,0x35)

C_NAVY='1F3864'; C_BLUE='2E75B6'; C_LIGHT='EEF4FB'
C_WHITE='FFFFFF'; C_ORANGE='C8590A'

MPL_NAVY='#1F3864'; MPL_BLUE='#2E75B6'; MPL_ACCENT='#00B0F0'
MPL_ORANGE='#C8590A'; MPL_GREEN='#538135'; MPL_GRAY='#888888'

def _norm(s):
    return unicodedata.normalize('NFD',str(s).lower()).encode('ascii','ignore').decode()

def set_cell_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex_color); tcPr.append(shd)

def add_heading(doc, text, level=1):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(10)
    p.paragraph_format.space_after=Pt(4)
    run=p.add_run(text); run.bold=True; run.font.name='Calibri'
    if level==1: run.font.size=Pt(13); run.font.color.rgb=NAVY
    elif level==2: run.font.size=Pt(11); run.font.color.rgb=BLUE
    else: run.font.size=Pt(10); run.font.color.rgb=GRAY
    return p

def add_body(doc, text, italic=False):
    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(4)
    run=p.add_run(text); run.font.name='Calibri'; run.font.size=Pt(10)
    run.font.italic=italic; run.font.color.rgb=RGBColor(0x33,0x33,0x33)
    return p

def fig_to_img(fig, width_cm=14):
    buf=BytesIO()
    fig.savefig(buf,format='png',dpi=150,bbox_inches='tight',facecolor='white')
    buf.seek(0); plt.close(fig)
    return buf, Cm(width_cm)

def add_section_header(doc, num, title):
    tbl=doc.add_table(rows=1,cols=1)
    tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
    c=tbl.rows[0].cells[0]; set_cell_bg(c,C_NAVY)
    p=c.paragraphs[0]
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    run=p.add_run(f'  {num}. {title.upper()}' if num else f'  {title.upper()}')
    run.font.name='Calibri'; run.font.size=Pt(11)
    run.font.bold=True; run.font.color.rgb=WHITE
    doc.add_paragraph()

def add_kpi_row(doc, kpis):
    tbl=doc.add_table(rows=2,cols=len(kpis))
    tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,(val,lbl) in enumerate(kpis):
        cv=tbl.rows[0].cells[j]; set_cell_bg(cv,C_LIGHT)
        p=cv.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(str(val)); run.font.name='Calibri'
        run.font.size=Pt(22); run.font.bold=True; run.font.color.rgb=BLUE
        cl=tbl.rows[1].cells[j]; set_cell_bg(cl,C_LIGHT)
        p2=cl.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run2=p2.add_run(lbl); run2.font.name='Calibri'
        run2.font.size=Pt(8.5); run2.font.color.rgb=GRAY
    doc.add_paragraph()

def _ax_style(ax):
    ax.set_facecolor('white')
    ax.spines[['top','right']].set_visible(False)
    ax.tick_params(labelsize=8)

def g_antes_despues(titulo, val1, val2, label1='Ingreso', label2='Seguimiento'):
    fig,ax=plt.subplots(figsize=(5,3))
    x=[0,1]; vals=[val1,val2]
    colors=[MPL_BLUE,MPL_ORANGE]
    bars=ax.bar([label1,label2],vals,color=colors,width=0.5)
    for b,v in zip(bars,vals):
        ax.text(b.get_x()+b.get_width()/2,b.get_height()+0.01*max(vals+[1]),
                f'{v}',ha='center',va='bottom',fontsize=11,fontweight='bold')
    _ax_style(ax); ax.yaxis.set_visible(False)
    ax.set_title(titulo,fontsize=10,fontweight='bold',color=MPL_NAVY)
    plt.tight_layout()
    return fig

def g_sust_comparativo(R):
    fig,ax=plt.subplots(figsize=(6,4))
    sust1=R.get('sust_dist_top1',{}); sust2=R.get('sust_dist_top2',{})
    todas=sorted(set(list(sust1.keys())+list(sust2.keys())))
    if not todas:
        ax.text(0.5,0.5,'Sin datos',ha='center',va='center',transform=ax.transAxes)
    else:
        x=np.arange(len(todas)); w=0.35
        v1=[sust1.get(s,0) for s in todas]
        v2=[sust2.get(s,0) for s in todas]
        ax.bar(x-w/2,v1,w,label='Ingreso',color=MPL_BLUE)
        ax.bar(x+w/2,v2,w,label='Seguimiento',color=MPL_ORANGE)
        ax.set_xticks(x); ax.set_xticklabels(todas,rotation=30,ha='right',fontsize=8)
        ax.legend(fontsize=8)
    _ax_style(ax)
    ax.set_title('Sustancia Principal — Ingreso vs Seguimiento',fontsize=10,fontweight='bold',color=MPL_NAVY)
    plt.tight_layout()
    return fig

def g_transgresion_comparativo(R):
    fig,ax=plt.subplots(figsize=(5,3))
    v1=R.get('pct_trans_top1',0); v2=R.get('pct_trans_top2',0)
    bars=ax.bar(['Ingreso','Seguimiento'],[v1,v2],color=[MPL_BLUE,MPL_ORANGE],width=0.5)
    for b,v in zip(bars,[v1,v2]):
        ax.text(b.get_x()+b.get_width()/2,b.get_height()+0.5,
                f'{v}%',ha='center',va='bottom',fontsize=10,fontweight='bold')
    _ax_style(ax); ax.yaxis.set_visible(False)
    ax.set_title('% con Transgresión — Ingreso vs Seguimiento',fontsize=10,fontweight='bold',color=MPL_NAVY)
    plt.tight_layout()
    return fig

def cargar_datos():
    df=pd.read_excel(INPUT_FILE,sheet_name=SHEET_NAME,header=1)
    df.columns=[str(c) for c in df.columns]

    seg=df[df.get('Tiene_TOP2','') == 'Sí'].copy() if 'Tiene_TOP2' in df.columns else \
        df[[c for c in df.columns if '_TOP2' in c and df[c].notna().any()]].dropna(how='all')

    # Intentar detectar columna Tiene_TOP2
    for col_tiene in ['Tiene_TOP2','tiene_top2']:
        if col_tiene in df.columns:
            seg=df[df[col_tiene]=='Sí'].copy().reset_index(drop=True)
            break

    N_total=len(df); N_seg=len(seg)
    R={'N_total':N_total,'N_seg':N_seg,'df':df,'seg':seg}

    def _col(cols, *keys):
        for c in cols:
            nc=_norm(c)
            if any(k in nc for k in keys): return c
        return None

    # Sustancias
    col_s1=next((c for c in df.columns if any(k in _norm(c) for k in
                 ['sustancia principal','cual considera','genera mas problemas'])
                 and '_TOP1' in c and 'RAW' not in c),None)
    col_s2=next((c for c in df.columns if any(k in _norm(c) for k in
                 ['sustancia principal','cual considera','genera mas problemas'])
                 and '_TOP2' in c and 'RAW' not in c),None)
    R['sust_dist_top1']=seg[col_s1].dropna().value_counts().head(6).to_dict() if col_s1 else {}
    R['sust_dist_top2']=seg[col_s2].dropna().value_counts().head(6).to_dict() if col_s2 else {}
    if R['sust_dist_top1']:
        R['sust_top1']=list(R['sust_dist_top1'].keys())[0]
        R['sust_top1_pct']=round(list(R['sust_dist_top1'].values())[0]/N_seg*100,1) if N_seg else 0
    else:
        R['sust_top1']=''; R['sust_top1_pct']=0

    # Transgresión
    col_t1=_col(seg.columns,'transgresion_top1','transgresion')
    for c in seg.columns:
        if 'transgresion' in _norm(c) and '_TOP1' in c: col_t1=c; break
    col_t2=None
    for c in seg.columns:
        if 'transgresion' in _norm(c) and '_TOP2' in c: col_t2=c; break
    R['n_trans_top1']=int((seg[col_t1]=='Sí').sum()) if col_t1 else 0
    R['n_trans_top2']=int((seg[col_t2]=='Sí').sum()) if col_t2 else 0
    R['pct_trans_top1']=round(R['n_trans_top1']/N_seg*100,1) if N_seg else 0
    R['pct_trans_top2']=round(R['n_trans_top2']/N_seg*100,1) if N_seg else 0

    # Salud
    col_sal1=_col(seg.columns,'estado de salud_top1','autopercep')
    for c in seg.columns:
        if 'salud' in _norm(c) and '_TOP1' in c: col_sal1=c; break
    col_sal2=None
    for c in seg.columns:
        if 'salud' in _norm(c) and '_TOP2' in c: col_sal2=c; break
    sal1=pd.to_numeric(seg[col_sal1],errors='coerce').dropna() if col_sal1 else pd.Series()
    sal2=pd.to_numeric(seg[col_sal2],errors='coerce').dropna() if col_sal2 else pd.Series()
    R['salud_media_top1']=round(float(sal1.mean()),1) if len(sal1) else 0
    R['salud_media_top2']=round(float(sal2.mean()),1) if len(sal2) else 0

    # Período
    global PERIODO
    col_f1=next((c for c in df.columns if 'fecha entrevista' in _norm(c) and '_TOP1' in c),None)
    if col_f1:
        fechas=pd.to_datetime(df[col_f1],errors='coerce').dropna()
        if len(fechas):
            MESES={1:'Enero',2:'Febrero',3:'Marzo',4:'Abril',5:'Mayo',6:'Junio',
                   7:'Julio',8:'Agosto',9:'Septiembre',10:'Octubre',11:'Noviembre',12:'Diciembre'}
            f1,f2=fechas.min(),fechas.max()
            if f1.year==f2.year and f1.month==f2.month:
                PERIODO=f'{MESES[f1.month]} {f1.year}'
            elif f1.year==f2.year:
                PERIODO=f'{MESES[f1.month]}–{MESES[f2.month]} {f1.year}'
            else:
                PERIODO=f'{MESES[f1.month]} {f1.year} – {MESES[f2.month]} {f2.year}'
    return R

def build_word(R):
    doc=Document()
    for sec in doc.sections:
        sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # Portada
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.rows[0].cells[0]; set_cell_bg(c,C_NAVY)
    for txt,sz,bold in [
        ('INFORME DE SEGUIMIENTO',18,True),
        ('Monitoreo de Resultados de Tratamiento — Instrumento TOP',11,False),
        ('Comparativo Ingreso vs Seguimiento (TOP1 – TOP2)',11,False),
        (NOMBRE_SERVICIO.upper(),14,True),
        (PERIODO,10,False),
    ]:
        p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(txt); run.font.name='Calibri'; run.font.size=Pt(sz)
        run.font.bold=bold; run.font.color.rgb=WHITE
    doc.add_paragraph()

    # Presentación
    add_section_header(doc,'','Presentación')
    add_body(doc,
        f'Este informe presenta los resultados comparativos entre el ingreso (TOP1) y el '
        f'seguimiento (TOP2) de {R["N_seg"]} personas, del total de {R["N_total"]} que '
        f'iniciaron tratamiento durante el período {PERIODO}.')
    doc.add_paragraph()

    kpis=[(R['N_total'],'Total ingresaron'),(R['N_seg'],'Con seguimiento')]
    pct=round(R['N_seg']/R['N_total']*100,1) if R['N_total'] else 0
    kpis.append((f'{pct}%','Tasa de seguimiento'))
    if R.get('sust_top1'): kpis.append((R['sust_top1'],'Sustancia principal'))
    add_kpi_row(doc,kpis)

    if R['N_seg'] == 0:
        add_body(doc,'No hay suficientes pacientes con seguimiento (TOP2) para generar el análisis comparativo.',italic=True)
        doc.save(OUTPUT_FILE)
        return

    # 1. Antecedentes
    add_section_header(doc,'1','Antecedentes del Seguimiento')
    add_body(doc,
        f'Del total de {R["N_total"]} personas que ingresaron a tratamiento, '
        f'{R["N_seg"]} ({pct}%) completaron el seguimiento al momento del análisis.')

    # 2. Consumo de Sustancias
    add_section_header(doc,'2','Consumo de Sustancias')
    add_heading(doc,'2.1. Sustancia Principal — Ingreso vs Seguimiento',level=2)
    if R.get('sust_dist_top1') or R.get('sust_dist_top2'):
        fig=g_sust_comparativo(R)
        buf,w=fig_to_img(fig,13)
        doc.add_picture(buf,width=w)
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        add_body(doc,
            f'Al ingreso, la sustancia principal más frecuente fue {R["sust_top1"]} '
            f'({R["sust_top1_pct"]}% de las personas con seguimiento).')

    # 3. Transgresión
    add_section_header(doc,'3','Transgresión a la Norma Social')
    add_heading(doc,'3.1. Transgresión — Ingreso vs Seguimiento',level=2)
    if R['N_seg'] > 0:
        fig=g_transgresion_comparativo(R)
        buf,w=fig_to_img(fig,11)
        doc.add_picture(buf,width=w)
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        cambio=round(R['pct_trans_top1']-R['pct_trans_top2'],1)
        direccion='disminuyó' if cambio>0 else 'aumentó' if cambio<0 else 'se mantuvo igual'
        add_body(doc,
            f'La proporción de personas con transgresión {direccion} de '
            f'{R["pct_trans_top1"]}% al ingreso a {R["pct_trans_top2"]}% en el seguimiento.')

    # 4. Salud
    add_section_header(doc,'4','Salud y Calidad de Vida')
    add_heading(doc,'4.1. Autopercepción del Estado de Salud',level=2)
    if R['salud_media_top1'] > 0 or R['salud_media_top2'] > 0:
        fig=g_antes_despues(
            'Estado de Salud (0–20)',
            R['salud_media_top1'],R['salud_media_top2'])
        buf,w=fig_to_img(fig,11)
        doc.add_picture(buf,width=w)
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        cambio_sal=round(R['salud_media_top2']-R['salud_media_top1'],1)
        dir_sal='mejoró' if cambio_sal>0 else 'disminuyó' if cambio_sal<0 else 'se mantuvo'
        add_body(doc,
            f'La autopercepción de salud {dir_sal} de {R["salud_media_top1"]} '
            f'puntos al ingreso a {R["salud_media_top2"]} puntos en el seguimiento '
            f'(escala 0–20).')

    # Pie
    doc.add_paragraph()
    p=doc.add_paragraph()
    run=p.add_run(f'Informe generado automáticamente por QALAT · Sistema de Monitoreo TOP · {PERIODO}')
    run.font.size=Pt(8); run.font.italic=True; run.font.color.rgb=GRAY
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_FILE)
    print(f'  ✓ Word generado: {OUTPUT_FILE}')


if __name__ == '__main__':
    R=cargar_datos()
    build_word(R)
